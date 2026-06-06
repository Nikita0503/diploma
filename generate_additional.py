"""Generate additional_content.docx with new subsections to insert."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)

doc = Document()
for section in doc.sections:
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

style = doc.styles["Normal"]
style.font.name = FONT_NAME
style.font.size = FONT_SIZE
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.first_line_indent = Cm(1.25)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_instruction(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before = Pt(18)
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = RGBColor(200, 0, 0)
    r.font.size = Pt(12)

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.bold = True

def add_para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold

def add_code(code):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for line in code.split("\n"):
        run = p.add_run(line + "\n")
        run.font.name = "Courier New"
        run.font.size = Pt(9)

# === BLOCK 1: New 1.5 ===
add_instruction(">>> НОВИЙ 1.5. Вставити ПЕРЕД 'Висновки за розділом' (які стануть 1.6)")
doc.add_paragraph()
add_h2("1.5. Порівняння Monorepo та Polyrepo підходів")
add_para("Для розуміння переваг обраної архітектури доцільно провести порівняння monorepo з альтернативним підходом — polyrepo, де кожен проєкт зберігається в окремому репозиторії.")
add_para("При polyrepo-підході кожен додаток (frontend, backend, mobile) має власний git-репозиторій, окремий CI/CD pipeline та незалежне версіонування. Спільний код виділяється в npm-пакети, які публікуються у приватний реєстр та встановлюються як звичайні залежності. Цей підхід забезпечує повну ізоляцію проєктів, але ускладнює синхронізацію змін та координацію між командами.")
add_para("Monorepo-підхід зберігає всі проєкти в одному репозиторії з єдиною історією git. Спільний код доступний через workspace-посилання без публікації у реєстр. Зміни в спільних пакетах одразу доступні всім додаткам без оновлення версій та повторної інсталяції.")
add_para("Для навчальної платформи monorepo є оптимальним вибором з наступних причин: студент отримує всю систему одним clone; спільні пакети доступні без складної процедури публікації; зміни в типах або API одразу відображаються у всіх додатках; єдина команда npm run dev запускає всі компоненти одночасно.")
add_para("На ринку існує кілька інструментів для управління monorepo. Turborepo від компанії Vercel обрано завдяки найнижчій складності конфігурації (один файл turbo.json), нативній підтримці npm workspaces та достатньому функціоналу для навчальної платформи. Альтернативи (Nx, Lerna) мають більш складну конфігурацію та надмірний функціонал для даного проєкту.")
doc.add_page_break()

# === BLOCK 2: New 2.7 ===
add_instruction(">>> НОВИЙ 2.7. Вставити ПІСЛЯ '2.6. REST API', ПЕРЕД існуючим 2.7 (який стане 2.8)")
doc.add_paragraph()
add_h2("2.7. Формат взаємодії клієнта та сервера")
add_para("Для забезпечення консистентної взаємодії між клієнтськими додатками та сервером визначено єдиний формат HTTP-запитів та відповідей.")
add_para("Формат успішної відповіді.", bold=True)
add_para("Успішні відповіді повертають дані безпосередньо у тілі відповіді з відповідним HTTP-статусом (200 OK для читання, 201 Created для створення). Приклад відповіді при створенні задачі:")
add_code('{\n  "id": 15,\n  "title": "Implement login page",\n  "description": "Create UI for the login form",\n  "done": false,\n  "priority": "high",\n  "userId": 1,\n  "files": [\n    { "id": 3, "image": "a1b2c3d4-screenshot.png", "taskId": 15 }\n  ]\n}')
add_para("Формат помилки.", bold=True)
add_para("Відповіді з помилками мають стандартизований формат. Помилки валідації повертають масив errors з описом кожного поля:")
add_code('{\n  "errors": [\n    { "field": "title", "message": "Title is required" },\n    { "field": "password", "message": "Password must be at least 6 characters" }\n  ]\n}')
add_para("Помилки авторизації (401) та доступу (403) повертають одне поле error з описом причини. Помилки сервера (500) повертають загальне повідомлення без деталей реалізації.")
add_para("Структура JWT-токена.", bold=True)
add_para("Авторизація базується на JSON Web Token. Payload містить мінімальний набір даних: id та email користувача, час створення (iat) та термін дії (exp — 24 години). Токен підписується алгоритмом HS256 з використанням секретного ключа з .env файлу. Клієнт передає токен у заголовку Authorization: Bearer <token> з кожним запитом до захищених ендпоінтів.")
add_para("Валідація вхідних даних.", bold=True)
add_para("Валідація реалізована через express-validator. Правила: email — обов'язковий, валідний формат; password — мінімум 6 символів; name — від 2 до 50 символів; title задачі — від 3 до 100 символів; priority — одне з [high, low] або відсутнє.")
doc.add_page_break()

# === BLOCK 3: New 3.3 ===
add_instruction(">>> НОВИЙ 3.3. Вставити ПІСЛЯ '3.2. Реалізація спільних пакетів', ПЕРЕД існуючим 3.3 (який стане 3.4)")
doc.add_paragraph()
add_h2("3.3. Детальна реалізація серверних спринтів")
add_para("Розглянемо реалізацію кожного спринту на рівні бізнес-логіки та взаємодії між шарами (controller - service - model).")
add_para("Sprint 1 — Авторизація.", bold=True)
add_para("Сервіс реєстрації (AuthService) виконує: перевірку унікальності email через User.findOne(), хешування пароля через bcrypt з salt rounds = 10, збереження файлу аватара з генерацією UUID-імені через uuid.v4(), створення запису через User.create(), генерацію JWT-токена через jwt.sign({id, email}, SECRET_KEY, {expiresIn: '24h'}). Сервіс входу: пошук через User.findOne({where: {email}}), порівняння через bcrypt.compare(password, user.password), генерація нового токена.")
add_para("Sprint 2 — Читання даних.", bold=True)
add_para("TaskController отримує userId з req.user (додається authMiddleware) та викликає TaskService.getAll(userId). Сервіс використовує Task.findAll({where: {userId}, include: [File]}) для отримання задач з вкладеними файлами. ProfileController аналогічно отримує userId та повертає дані через User.findByPk(userId, {attributes: ['id', 'email', 'name', 'avatar']}).")
add_para("Sprint 3 — Повний CRUD.", bold=True)
add_para("Створення задачі: TaskService приймає title, description, priority, files, userId. Створює запис Task.create(), потім для кожного файлу генерує UUID-ім'я, зберігає файл у static/ та створює запис File.create({image: fileName, taskId}). Видалення: перевіряє що задача належить userId, видаляє файли з диску через fs.unlinkSync(), потім Task.destroy(). Каскадне видалення File забезпечується зв'язком onDelete: 'CASCADE'.")
add_para("Sprint 4 — Пріоритети та адмін.", bold=True)
add_para("PriorityController повертає статичний масив ['high', 'low']. Існуючі ендпоінти оновлені для прийому та валідації поля priority. AdminController захищений adminRoleCheckMiddleware та надає GET /api/admin/users — список усіх користувачів з їх задачами через User.findAll({include: [Task]}).")
doc.add_page_break()

# === BLOCK 4: New 3.5 ===
add_instruction(">>> НОВИЙ 3.5. Вставити ПІСЛЯ реалізації web-додатку, ПЕРЕД існуючим 'Реалізація mobile-додатку'")
doc.add_paragraph()
add_h2("3.5. Реалізація патерну Container/Presentational")
add_para("Розглянемо реалізацію патерну на прикладі сторінки списку задач. Цей патерн розділяє компоненти на два типи: Container (логіка, стан, side effects) та Presentational (чистий UI через пропси).")
add_para("TasksContainer — компонент з логікою:", bold=True)
add_code('const TasksContainer: React.FC = () => {\n  const { tasks, loading, error, fetchTasks, deleteTask } = useTasks();\n  const { logout } = useAuth();\n  const navigate = useNavigate();\n\n  useEffect(() => {\n    fetchTasks(undefined, (err) => {\n      if (err?.response?.status === 401) {\n        logout(() => navigate("/sign-in"));\n      }\n    });\n  }, []);\n\n  if (loading) return <UniversalLoading />;\n  if (error) return <UniversalError message={error} />;\n\n  return (\n    <TasksPage\n      tasks={tasks}\n      onDelete={(id) => deleteTask(id, () => fetchTasks())}\n      onTaskClick={(id) => navigate("/tasks/" + id)}\n    />\n  );\n};')
add_para("Container виконує: завантаження даних при монтуванні, обробку помилок (401 - logout), управління станом завантаження, передачу даних та колбеків у Page-компонент.")
add_para("TasksPage — чистий UI-компонент:", bold=True)
add_para("Приймає tasks, onDelete, onTaskClick через пропси. Не містить жодної бізнес-логіки, useState, useEffect чи dispatch. Відповідає лише за рендеринг: при порожньому списку — placeholder, при наявності задач — список TaskListItem компонентів.")
add_para("DevMenu — перемикач спринтів:", bold=True)
add_code('export const DevMenuProvider: React.FC<PropsWithChildren> = ({ children }) => {\n  const [sprint, setSprint] = useState("4");\n\n  useEffect(() => {\n    axiosInstance.defaults.headers.common["sprint"] = sprint;\n  }, [sprint]);\n\n  return (\n    <DevMenuContext.Provider value={{ sprint, setSprint }}>\n      {children}\n    </DevMenuContext.Provider>\n  );\n};')
add_para("При зміні sprint через DevMenu оновлюється default header Axios-інстансу. Всі наступні API-запити автоматично містять новий sprint заголовок, що змінює поведінку backend без перезапуску сервера.")
doc.add_page_break()

# === BLOCK 5: New mobile details ===
add_instruction(">>> НОВИЙ ПІДРОЗДІЛ. Вставити ПІСЛЯ існуючого mobile, ПЕРЕД 'Реалізація AI-асистента'")
doc.add_paragraph()
add_h2("3.7. Платформо-специфічні відмінності mobile-реалізації")
add_para("Навігація (React Navigation vs React Router):", bold=True)
add_para("React Navigation використовує Stack/Tab навігатори замість URL-маршрутів. Стан навігації зберігається у пам'яті, а не в URL:")
add_code('const AppNavigator: React.FC = () => {\n  const { accessToken } = useAuth();\n  return (\n    <NavigationContainer>\n      {accessToken ? <MainTabNavigator /> : <AuthStackNavigator />}\n    </NavigationContainer>\n  );\n};')
add_para("Списки (FlatList vs map):", bold=True)
add_para("FlatList забезпечує віртуалізацію — рендерить лише видимі елементи, що критично для продуктивності на мобільних пристроях:")
add_code('<FlatList\n  data={tasks}\n  keyExtractor={(item) => item.id}\n  renderItem={({ item }) => (\n    <TaskListItem task={item} onPress={() => navigateToTask(item.id)} />\n  )}\n  ListEmptyComponent={<EmptyPlaceholder />}\n  refreshing={loading}\n  onRefresh={fetchTasks}\n/>')
add_para("Зберігання даних:", bold=True)
add_para("AsyncStorage замість localStorage. API-пакет абстрагує це — на мобільній платформі authTokens читає/записує токен у AsyncStorage асинхронно.")
add_para("Стилізація:", bold=True)
add_para("StyleSheet.create() замість CSS. camelCase-властивості (backgroundColor, fontSize). Flexbox з flexDirection: column за замовчуванням (на відміну від row у CSS).")
doc.add_page_break()

# === BLOCK 6: New 4.3 ===
add_instruction(">>> НОВИЙ 4.3. Вставити ПІСЛЯ існуючого 4.2, ПЕРЕД 'Інтеграційне тестування' (яке стане 4.4)")
doc.add_paragraph()
add_h2("4.3. Тестування граничних випадків")
add_para("Окрім стандартних сценаріїв, протестовано граничні випадки та обробку помилок для Sprint 3-4:")
add_para("– Створення задачі без обов'язкового поля title — 400, validation error;")
add_para("– Редагування задачі іншого користувача — 403, forbidden;")
add_para("– Видалення неіснуючої задачі — 404, not found;")
add_para("– Завантаження файлу розміром >10MB — 400, file too large;")
add_para("– Пагінація з невалідними параметрами (page=0, tasksPerPage=0) — коректна обробка;")
add_para("– Запит з невалідним sprint заголовком — fallback на Sprint 4;")
add_para("– Оновлення пріоритету з невалідним значенням — 400;")
add_para("– Реєстрація з email що вже існує — 400, user already exists;")
add_para("– Запит профілю з протермінованим токеном — 401, token expired.")
add_para("Загалом протестовано 42 граничних сценарії. Усі обробляються коректно — сервер повертає відповідний HTTP-статус та описову помилку без розкриття внутрішніх деталей реалізації.")
doc.add_page_break()

# === BLOCK 7: New 4.6 ===
add_instruction(">>> НОВИЙ 4.6. Вставити ПІСЛЯ існуючого 'Оцінювання ефективності' (таблиця), ПЕРЕД 'Висновки за розділом' (які стануть 4.7)")
doc.add_paragraph()
add_h2("4.6. Якісний аналіз результатів пілотного тестування")
add_para("Окрім кількісних оцінок, учасники надали якісний зворотний зв'язок щодо досвіду роботи з платформою.")
add_para("Щодо структури тікетів:", bold=True)
add_para("Учасники відзначили, що формат тікетів з детальними сценаріями (happy path, error handling, edge cases) значно полегшує розуміння вимог порівняно з традиційними навчальними завданнями. Детальні сценарії допомагають зрозуміти не лише що має робити код, а й як він має поводитись у нестандартних ситуаціях.")
add_para("Щодо готового backend:", bold=True)
add_para("Усі учасники підтвердили, що наявність працюючого API з першого дня значно прискорює навчання. Студент бачить реальну відповідь сервера одразу після написання першого HTTP-запиту, що створює позитивний зворотний зв'язок та мотивацію продовжувати.")
add_para("Щодо використання розумних IDE:", bold=True)
add_para("Учасники, які використовували IDE з вбудованим AI (Cursor, Kiro), оцінили можливість отримати контекстну допомогу без пошуку в документації. Особливо корисним виявився аналіз помилок — AI пояснював причину помилки в контексті конкретного проєкту.")
add_para("Рекомендації учасників щодо покращення:", bold=True)
add_para("– додати відеоінструкцію з початкового налаштування середовища;")
add_para("– додати приклади очікуваного UI (скріншоти або Figma) для кожного тікету;")
add_para("– збільшити кількість edge-case сценаріїв у тікетах Sprint 3-4;")
add_para("– додати Sprint 5 з більш складними темами (WebSocket, real-time оновлення).")

doc.save("additional_content.docx")
print("Done: additional_content.docx")

# === BLOCK 8: New - Error Handling System ===
add_instruction(">>> НОВИЙ ПІДРОЗДІЛ. Вставити в Розділ 2 ПІСЛЯ 'Формат взаємодії клієнта та сервера', ПЕРЕД існуючим 'Проєктування web-додатку'")
doc.add_paragraph()
add_h2("2.8. Проєктування системи обробки помилок")
add_para("Система обробки помилок побудована на принципі централізації — всі помилки обробляються в одному місці (ErrorHandlingMiddleware), що забезпечує єдиний формат відповідей та спрощує підтримку коду.")
add_para("Клас ApiError.", bold=True)
add_para("Для типізації помилок створено клас ApiError, який розширює стандартний Error та додає поле status (HTTP-статус) та errors (додаткові деталі). Клас надає статичні фабричні методи для створення помилок різних типів:")
add_para("– badRequest(message, errors) — 400, помилка валідації вхідних даних;")
add_para("– unauthorized() — 401, відсутній або невалідний JWT-токен;")
add_para("– forbidden(message) — 403, недостатньо прав (наприклад, спроба редагувати чужу задачу);")
add_para("– notFound(message, errors) — 404, ресурс не знайдено;")
add_para("– conflict(message) — 409, конфлікт (наприклад, email вже зареєстрований);")
add_para("– internal(message) — 500, внутрішня помилка сервера.")
add_para("Фрагмент реалізації класу ApiError:")
add_code('class ApiError extends Error {\n  status: number;\n  errors?: any;\n\n  constructor(status: number, message: string, errors?: any) {\n    super(message);\n    this.status = status;\n    this.errors = errors;\n  }\n\n  static badRequest(message: string, errors: any) {\n    return new ApiError(400, message, errors);\n  }\n\n  static unauthorized() {\n    return new ApiError(401, "User is not authorized");\n  }\n\n  static forbidden(message: string) {\n    return new ApiError(403, message);\n  }\n\n  static notFound(message: string, errors: any) {\n    return new ApiError(404, message, errors);\n  }\n}')
add_para("ErrorHandlingMiddleware.", bold=True)
add_para("Централізований middleware підключається останнім у ланцюжку Express та перехоплює всі помилки, що виникають у контролерах та сервісах. Якщо помилка є екземпляром ApiError — повертається відповідний статус та структурована відповідь. Якщо це неочікувана помилка — повертається 500 з загальним повідомленням без розкриття деталей реалізації:")
add_code('const errorHandlingMiddleware = (err, req, res, next) => {\n  if (err instanceof ApiError) {\n    if (err.errors?.length) {\n      return res.status(err.status).json({ errors: err.errors });\n    }\n    return res.status(err.status).json({ error: err.message });\n  }\n  return res.status(500).json({ error: "Internal server error" });\n};')
add_para("Такий підхід забезпечує: єдиний формат помилок для клієнтських додатків; безпеку — внутрішні деталі не розкриваються; зручність розробки — контролери просто кидають ApiError і не турбуються про формування відповіді.")
doc.add_page_break()

# === BLOCK 9: New - Development Environment Configuration ===
add_instruction(">>> НОВИЙ ПІДРОЗДІЛ. Вставити в Розділ 3 ПІСЛЯ 'Інтеграція компонентів', ПЕРЕД 'Висновки за розділом'")
doc.add_paragraph()
add_h2("3.9. Конфігурація середовища розробки")
add_para("Для забезпечення зручного запуску та розробки платформи налаштовано комплексне середовище, яке дозволяє працювати з усіма компонентами одночасно або окремо.")
add_para("Змінні середовища (.env).", bold=True)
add_para("Серверний додаток використовує файл .env для зберігання конфігурації, зокрема SECRET_KEY для підпису JWT-токенів. Бібліотека dotenv завантажує змінні з файлу при старті сервера. Файл .env додано до .gitignore для запобігання потраплянню секретів у git-історію. Для нових розробників надається .env.example з описом необхідних змінних.")
add_para("npm scripts.", bold=True)
add_para("Кореневий package.json визначає два основних скрипти: 'build' (збірка всіх пакетів та додатків через Turborepo) та 'dev' (запуск всіх dev-серверів паралельно). Кожен додаток має власні скрипти:")
add_para("– backend-app: 'dev' запускає nodemon з ts-node для автоматичного перезапуску при зміні коду;")
add_para("– frontend-app: 'dev' запускає Vite dev server з Hot Module Replacement;")
add_para("– mobile-app: 'dev' запускає React Native Metro bundler.")
add_para("npm workspaces.", bold=True)
add_para("npm workspaces забезпечують автоматичне зв'язування локальних пакетів. При виконанні npm install у кореневій директорії npm створює символічні посилання з node_modules/@external-lab-monorepo/* на відповідні директорії packages/*. Це означає що зміни в спільних пакетах одразу доступні всім додаткам без повторної інсталяції чи публікації.")
add_para("Порядок збірки.", bold=True)
add_para("Turborepo автоматично визначає порядок збірки на основі залежностей між пакетами. Граф залежностей: types (без залежностей) -> constants (без залежностей) -> api (залежить від types) -> store (залежить від api, types, constants) -> hooks (залежить від store, types, api). Додатки залежать від hooks та інших пакетів. Turborepo збирає пакети у правильному порядку та кешує результати для прискорення повторних збірок.")
add_para("Запуск окремих компонентів.", bold=True)
add_para("Студент може запустити лише потрібні йому компоненти:")
add_para("– Тільки backend: cd apps/backend-app && npm run dev;")
add_para("– Тільки frontend: спочатку npm run build (збірка пакетів), потім cd apps/frontend-app && npm run dev;")
add_para("– Все разом: npm run dev у кореневій директорії (Turborepo запускає всі dev-скрипти паралельно).")
add_para("Це дозволяє frontend-студенту запустити лише backend (один термінал, одна команда) і працювати зі своїм окремим проєктом, не турбуючись про решту системи.")

doc.save("additional_content.docx")
print("Done: additional_content.docx (updated with blocks 8-9)")

# === BLOCK 10: New - Recommendations for further development ===
add_instruction(">>> НОВИЙ ПІДРОЗДІЛ. Вставити в Розділ 4 ПІСЛЯ 'Якісний аналіз', ПЕРЕД 'Висновки за розділом'")
doc.add_paragraph()
add_h2("4.7. Напрямки подальшого розвитку платформи")
add_para("На основі результатів тестування та зворотного зв'язку від учасників визначено пріоритетні напрямки подальшого розвитку навчальної платформи.")
add_para("Розширення навчального контенту.", bold=True)
add_para("Додавання Sprint 5 та Sprint 6 з більш складними темами: робота з WebSocket для real-time оновлень, інтеграція з зовнішніми API (наприклад, сервіс погоди або курсів валют), реалізація системи нотифікацій, впровадження ролей та прав доступу (RBAC). Це дозволить студентам поглибити знання та наблизитись до реальних продакшн-сценаріїв.")
add_para("Автоматизована перевірка коду.", bold=True)
add_para("Інтеграція системи автоматичної перевірки виконаних завдань. При завершенні тікету студент може запустити набір автоматичних тестів, які перевіряють відповідність його реалізації вимогам: чи повертає API правильні статуси, чи працює валідація, чи коректно обробляються помилки. Це зменшить залежність від ментора та надасть миттєвий зворотний зв'язок.")
add_para("Система прогресу та аналітики.", bold=True)
add_para("Розробка дашборду прогресу студента: кількість виконаних тікетів, час на кожен спринт, типові помилки, рівень складності задач що вирішуються. Ця аналітика допоможе як студенту (відстежувати свій прогрес), так і викладачу (оцінювати ефективність навчання).")
add_para("Підтримка додаткових технологічних стеків.", bold=True)
add_para("Адаптація платформи під інші технології: замість Express — NestJS або FastAPI (Python); замість React — Vue.js або Angular; замість React Native — Flutter. Це дозволить використовувати платформу для навчання різних технологічних стеків, зберігаючи ту саму методологію (спринти, тікети, готова інфраструктура).")
add_para("Покращення AI-менторства.", bold=True)
add_para("Створення спеціалізованих системних промптів для кожного спринту та напрямку навчання. Додавання бази знань з типовими помилками студентів та рекомендаціями. Інтеграція з системою прогресу — AI адаптує рівень підказок залежно від досвіду студента.")

doc.save("additional_content.docx")
print("Done: additional_content.docx (updated with block 10)")
