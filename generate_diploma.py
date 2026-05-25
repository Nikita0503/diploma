"""
Генератор дипломної роботи у форматі .docx
Тема: Розроблення навчальної платформи для набуття практичних навичок
       full-stack розробки на основі Monorepo-архітектури
Автор: Ключак Ольга Андріївна
"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# --- Configuration ---
OUTPUT_FILE = "diploma.docx"
FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(14)
LINE_SPACING = 1.5
FIRST_LINE_INDENT = Cm(1.25)
MARGIN_LEFT = Cm(3.0)
MARGIN_RIGHT = Cm(1.0)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)


def create_document():
    """Create a new document with proper page setup."""
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = MARGIN_LEFT
        section.right_margin = MARGIN_RIGHT
        section.top_margin = MARGIN_TOP
        section.bottom_margin = MARGIN_BOTTOM

    setup_styles(doc)
    return doc


def setup_styles(doc):
    """Configure document styles."""
    # Normal style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.first_line_indent = FIRST_LINE_INDENT
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set font for East Asian and Complex Script
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    rFonts.set(qn("w:ascii"), FONT_NAME)
    rFonts.set(qn("w:hAnsi"), FONT_NAME)
    rFonts.set(qn("w:cs"), FONT_NAME)

    # Heading 1 - Chapter titles (РОЗДІЛ 1. ...)
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
    else:
        h1 = doc.styles.add_style("Heading 1", WD_STYLE_TYPE.PARAGRAPH)
    h1.font.name = FONT_NAME
    h1.font.size = FONT_SIZE
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Cm(0)
    h1.paragraph_format.line_spacing = LINE_SPACING
    h1.paragraph_format.keep_with_next = True

    # Heading 2 - Subsection titles (1.1. ...)
    if "Heading 2" in doc.styles:
        h2 = doc.styles["Heading 2"]
    else:
        h2 = doc.styles.add_style("Heading 2", WD_STYLE_TYPE.PARAGRAPH)
    h2.font.name = FONT_NAME
    h2.font.size = FONT_SIZE
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    h2.paragraph_format.line_spacing = LINE_SPACING
    h2.paragraph_format.keep_with_next = True


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


def add_chapter_title(doc, text):
    """Add a chapter title (РОЗДІЛ X. НАЗВА)."""
    para = doc.add_paragraph(text.upper(), style="Heading 1")
    return para


def add_subsection_title(doc, text):
    """Add a subsection title (X.X. Назва)."""
    para = doc.add_paragraph(text, style="Heading 2")
    return para


def add_paragraph(doc, text, bold=False, italic=False, indent=True):
    """Add a normal paragraph."""
    para = doc.add_paragraph()
    if not indent:
        para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    return para


def add_centered_paragraph(doc, text, bold=False, size=None):
    """Add a centered paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run(text)
    run.bold = bold
    if size:
        run.font.size = size
    return para


def add_list_item(doc, text, level=0):
    """Add a list item with dash."""
    prefix = "– " if level == 0 else "  – "
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = para.add_run(prefix + text)
    return para


def add_numbered_item(doc, number, text):
    """Add a numbered list item."""
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
    run = para.add_run(f"{number}. {text}")
    return para


def add_image(doc, image_path, width=None):
    """Add an image centered."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run()
    if width:
        run.add_picture(image_path, width=width)
    else:
        run.add_picture(image_path, width=Cm(14))
    return para


def add_image_caption(doc, text):
    """Add a figure caption."""
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.first_line_indent = Cm(0)
    para.paragraph_format.space_before = Pt(6)
    run = para.add_run(text)
    run.italic = True
    run.font.size = Pt(12)
    return para


# ============================================================
# CONTENT SECTIONS
# ============================================================

def write_annotation(doc):
    """Write annotation in Ukrainian."""
    add_chapter_title(doc, "АНОТАЦІЯ")
    add_paragraph(doc, (
        "Бакалаврська кваліфікаційна робота на тему: «Розроблення навчальної "
        "платформи для набуття практичних навичок full-stack розробки на основі "
        "Monorepo-архітектури» присвячена розробці програмного комплексу, який "
        "забезпечує ефективне навчання студентів практичним навичкам веб- та "
        "мобільної розробки шляхом роботи з реальним проєктом у monorepo-середовищі."
    ))
    add_paragraph(doc, (
        "Актуальність дослідження зумовлена невідповідністю між теоретичними "
        "знаннями, які отримують студенти IT-спеціальностей, та практичними "
        "навичками, що вимагаються на ринку праці. Існуючі навчальні платформи "
        "переважно зосереджені на ізольованих вправах і не надають досвіду роботи "
        "з повноцінною багатокомпонентною системою."
    ))
    add_paragraph(doc, (
        "У процесі виконання роботи було спроєктовано та реалізовано навчальну "
        "платформу, що містить три додатки (серверний, веб-клієнт, мобільний "
        "клієнт) та п'ять спільних пакетів, об'єднаних у monorepo-архітектурі "
        "з використанням Turborepo. Платформа реалізує спринтову модель навчання "
        "з поступовим ускладненням завдань та інтегрований AI-асистент для "
        "супроводу навчального процесу."
    ))

    add_paragraph(doc, (
        "Наукова новизна полягає у поєднанні monorepo-архітектури як навчального "
        "інструменту зі спринтовою методологією та AI-менторством, що створює "
        "унікальне середовище для самостійного набуття практичних навичок "
        "full-stack розробки."
    ))
    add_paragraph(doc, (
        "Практичне значення роботи полягає у створенні готового до використання "
        "навчального комплексу, який може бути застосований у вищих навчальних "
        "закладах, на курсах підвищення кваліфікації та для самостійного навчання."
    ))
    add_paragraph(doc, (
        "Ключові слова: навчальна платформа, monorepo, full-stack розробка, "
        "React, React Native, Express, TypeScript, Turborepo, спринтова "
        "методологія, AI-асистент, shared packages."
    ))


def write_annotation_en(doc):
    """Write annotation in English."""
    add_page_break(doc)
    add_chapter_title(doc, "ANNOTATION")
    add_paragraph(doc, (
        "Bachelor's thesis topic: \"Development of an Educational Platform for "
        "Acquiring Practical Full-Stack Development Skills Based on Monorepo "
        "Architecture\"."
    ))
    add_paragraph(doc, (
        "The thesis is dedicated to the development of a software system that "
        "provides effective training of students in practical web and mobile "
        "development skills by working with a real project in a monorepo "
        "environment."
    ))
    add_paragraph(doc, (
        "The relevance of the study is driven by the gap between theoretical "
        "knowledge obtained by IT students and practical skills required in the "
        "job market. Existing educational platforms are mostly focused on isolated "
        "exercises and do not provide experience working with a full-fledged "
        "multi-component system."
    ))

    add_paragraph(doc, (
        "During the work, an educational platform was designed and implemented, "
        "containing three applications (server, web client, mobile client) and "
        "five shared packages, united in a monorepo architecture using Turborepo. "
        "The platform implements a sprint-based learning model with progressive "
        "task complexity and an integrated AI assistant to support the learning "
        "process."
    ))
    add_paragraph(doc, (
        "Keywords: educational platform, monorepo, full-stack development, "
        "React, React Native, Express, TypeScript, Turborepo, sprint methodology, "
        "AI assistant, shared packages."
    ))


def write_abbreviations(doc):
    """Write list of abbreviations."""
    add_page_break(doc)
    add_chapter_title(doc, "ПЕРЕЛІК УМОВНИХ ПОЗНАЧЕНЬ ТА СКОРОЧЕНЬ")
    abbreviations = [
        ("API", "Application Programming Interface — програмний інтерфейс взаємодії між компонентами системи"),
        ("CRUD", "Create, Read, Update, Delete — базові операції з даними"),
        ("CSS", "Cascading Style Sheets — каскадні таблиці стилів"),
        ("DOM", "Document Object Model — об'єктна модель документа"),
        ("ER", "Entity-Relationship — модель «сутність-зв'язок»"),
        ("HTTP", "HyperText Transfer Protocol — протокол передачі гіпертексту"),
        ("JSON", "JavaScript Object Notation — формат обміну даними"),
        ("JWT", "JSON Web Token — стандарт токенів автентифікації"),
        ("LLM", "Large Language Model — велика мовна модель"),
        ("ORM", "Object-Relational Mapping — об'єктно-реляційне відображення"),
        ("REST", "Representational State Transfer — архітектурний стиль API"),
        ("SDK", "Software Development Kit — набір інструментів розробки"),
        ("SPA", "Single Page Application — односторінковий додаток"),
        ("SQL", "Structured Query Language — мова структурованих запитів"),
        ("UI", "User Interface — інтерфейс користувача"),
        ("UX", "User Experience — досвід користувача"),
    ]
    for abbr, desc in abbreviations:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        run = para.add_run(f"{abbr}")
        run.bold = True
        para.add_run(f" — {desc}.")


def write_introduction(doc):
    """Write ВСТУП section."""
    add_page_break(doc)
    add_chapter_title(doc, "ВСТУП")

    add_paragraph(doc, (
        "Сучасний ринок праці у сфері інформаційних технологій висуває високі "
        "вимоги до практичних навичок розробників програмного забезпечення. "
        "Роботодавці очікують від кандидатів не лише знання окремих мов "
        "програмування чи фреймворків, а й досвід роботи з повноцінними "
        "багатокомпонентними системами, розуміння архітектурних патернів та "
        "вміння працювати з кодовою базою, що складається з кількох взаємопов'язаних "
        "додатків. Проте традиційна система вищої освіти та більшість існуючих "
        "онлайн-платформ не забезпечують такого досвіду, зосереджуючись на "
        "ізольованих вправах та теоретичних завданнях."
    ))

    add_paragraph(doc, (
        "Особливо гострою є проблема набуття навичок full-stack розробки — "
        "одночасного розуміння серверної, клієнтської та мобільної частин "
        "програмного продукту. Студент, який вивчає frontend-розробку, рідко "
        "має можливість працювати з реальним backend-сервером, а той, хто "
        "вивчає backend, не бачить як його API використовується реальними "
        "клієнтськими додатками. Це призводить до фрагментарного розуміння "
        "процесу розробки та неготовності до роботи в команді."
    ))

    add_paragraph(doc, (
        "Monorepo-архітектура — підхід до організації коду, при якому кілька "
        "пов'язаних проєктів зберігаються в одному репозиторії — набуває "
        "дедалі більшої популярності у великих технологічних компаніях (Google, "
        "Meta, Microsoft). Цей підхід дозволяє ефективно перевикористовувати "
        "код між різними додатками через спільні пакети, забезпечує єдину "
        "систему збірки та спрощує координацію між командами. Використання "
        "monorepo як основи навчальної платформи дозволяє студенту побачити "
        "повну картину системи та зрозуміти взаємозв'язки між її компонентами."
    ))

    add_paragraph(doc, (
        "Актуальність теми бакалаврської роботи зумовлена кількома факторами. "
        "По-перше, за даними Stack Overflow Developer Survey 2024, понад 70% "
        "розробників вважають практичний досвід найважливішим фактором при "
        "працевлаштуванні, тоді як лише 30% задоволені якістю практичної "
        "підготовки у навчальних закладах. По-друге, monorepo-архітектура "
        "стає індустріальним стандартом для організації складних проєктів, "
        "проте навчальних матеріалів, що демонструють цей підхід на практиці, "
        "критично мало. По-третє, розвиток технологій штучного інтелекту "
        "відкриває нові можливості для персоналізованого навчання через "
        "AI-асистентів, які можуть виступати в ролі ментора."
    ))

    add_paragraph(doc, (
        "Мета роботи полягає у розробленні навчальної платформи на основі "
        "monorepo-архітектури, яка забезпечує набуття практичних навичок "
        "full-stack розробки через роботу з реальним багатокомпонентним "
        "проєктом та поспринтову систему завдань з підтримкою AI-асистента."
    ))

    add_paragraph(doc, "Завдання дослідження включають:", bold=False)
    tasks = [
        "Провести аналіз існуючих навчальних платформ для розробників та виявити їх обмеження.",
        "Обґрунтувати вибір monorepo-архітектури як основи навчальної платформи.",
        "Спроєктувати архітектуру системи, що включає серверний додаток, веб-клієнт, мобільний клієнт та спільні пакети.",
        "Реалізувати навчальну платформу з використанням сучасних технологій (TypeScript, React, React Native, Express, Turborepo).",
        "Розробити спринтову систему навчальних завдань з поступовим ускладненням.",
        "Інтегрувати AI-асистент для супроводу навчального процесу.",
        "Провести тестування платформи та оцінити її ефективність.",
    ]
    for i, task in enumerate(tasks, 1):
        add_numbered_item(doc, i, task)

    add_paragraph(doc, (
        "Об'єкт дослідження — процес навчання практичним навичкам full-stack "
        "розробки програмного забезпечення."
    ))

    add_paragraph(doc, (
        "Предмет дослідження — методи та засоби побудови навчальної платформи "
        "на основі monorepo-архітектури з інтегрованим AI-асистентом."
    ))

    add_paragraph(doc, "Методи дослідження включають:")
    methods = [
        "аналіз літературних джерел та існуючих рішень у сфері навчання розробці програмного забезпечення;",
        "порівняльний аналіз архітектурних підходів до організації коду;",
        "системний підхід до проєктування багатокомпонентної системи;",
        "методи програмної інженерії для реалізації платформи;",
        "експериментальний метод для тестування та верифікації результатів.",
    ]
    for method in methods:
        add_list_item(doc, method)

    add_paragraph(doc, (
        "Наукова новизна одержаних результатів полягає у розробленні концепції "
        "навчальної платформи, що поєднує monorepo-архітектуру як навчальний "
        "інструмент зі спринтовою методологією поступового ускладнення завдань "
        "та AI-менторством. На відміну від існуючих рішень, платформа надає "
        "студенту повноцінне робоче середовище з готовими компонентами системи, "
        "що дозволяє зосередитись на вивченні обраного напрямку, одночасно "
        "розуміючи контекст всієї системи."
    ))

    add_paragraph(doc, (
        "Практична цінність роботи полягає у створенні готового до використання "
        "навчального комплексу, який включає: працюючий серверний додаток з "
        "REST API, веб-клієнт на React, мобільний клієнт на React Native, "
        "набір спільних пакетів, документовану систему навчальних завдань "
        "та інтегрований AI-асистент. Платформа може бути використана у "
        "вищих навчальних закладах, на курсах підвищення кваліфікації та "
        "для самостійного навчання."
    ))


def write_chapter1(doc):
    """Write РОЗДІЛ 1 - Аналіз предметної області."""
    add_page_break(doc)
    add_chapter_title(doc, "РОЗДІЛ 1. АНАЛІЗ ПРЕДМЕТНОЇ ОБЛАСТІ")

    # 1.1
    add_subsection_title(doc, "1.1. Проблема набуття практичних навичок full-stack розробки")

    add_paragraph(doc, (
        "Сучасна індустрія розробки програмного забезпечення характеризується "
        "високим темпом розвитку технологій та зростаючими вимогами до "
        "кваліфікації спеціалістів. Full-stack розробник — це фахівець, який "
        "володіє навичками створення як серверної (backend), так і клієнтської "
        "(frontend) частин веб-додатків, а також, дедалі частіше, мобільних "
        "додатків. За даними LinkedIn Workforce Report 2024, попит на "
        "full-stack розробників зріс на 35% порівняно з попереднім роком, "
        "що свідчить про актуальність цього напрямку."
    ))

    add_paragraph(doc, (
        "Основна проблема підготовки full-stack розробників полягає у "
        "фрагментарності навчального процесу. Типова навчальна програма "
        "розділяє вивчення frontend та backend на окремі курси, які рідко "
        "інтегруються між собою. Студент вивчає React окремо від Express, "
        "базу даних окремо від API, мобільну розробку окремо від серверної "
        "частини. В результаті випускник має набір ізольованих знань, але "
        "не розуміє як ці компоненти працюють разом у реальному проєкті."
    ))

    add_paragraph(doc, (
        "Додатковою проблемою є відсутність досвіду роботи з реальною "
        "кодовою базою. У навчальних проєктах студент зазвичай пише весь "
        "код самостійно з нуля, що не відповідає реальним умовам роботи, "
        "де розробник приєднується до існуючого проєкту, працює з чужим "
        "кодом та інтегрується з компонентами, написаними іншими членами "
        "команди. Навичка читання та розуміння чужого коду, роботи з "
        "контрактами API та спільними бібліотеками є критично важливою, "
        "але практично не розвивається у традиційному навчальному процесі."
    ))

    add_paragraph(doc, (
        "Ще одним аспектом проблеми є відсутність поступовості у навчанні. "
        "Більшість курсів або занадто прості (todo-додатки), або одразу "
        "занадто складні (повноцінні SaaS-платформи). Студенту бракує "
        "проміжного етапу — проєкту, який починається з простих задач "
        "(авторизація) і поступово ускладнюється (CRUD-операції, робота "
        "з файлами, пагінація, пріоритети), імітуючи реальний процес "
        "розробки продукту спринтами."
    ))

    add_paragraph(doc, (
        "Таким чином, існує потреба у навчальній платформі, яка б: "
        "надавала досвід роботи з повноцінною багатокомпонентною системою; "
        "дозволяла зосередитись на одному напрямку (frontend, backend або "
        "mobile), маючи решту компонентів готовими; забезпечувала поступове "
        "ускладнення завдань; та надавала підтримку у вигляді документації "
        "та AI-менторства."
    ))

    # 1.2
    add_subsection_title(doc, "1.2. Огляд існуючих навчальних платформ та їх обмеження")

    add_paragraph(doc, (
        "На сучасному ринку існує значна кількість платформ для навчання "
        "програмуванню. Розглянемо найбільш популярні з них та проаналізуємо "
        "їх переваги та обмеження з точки зору набуття практичних навичок "
        "full-stack розробки."
    ))

    add_paragraph(doc, (
        "freeCodeCamp — безкоштовна платформа з інтерактивними завданнями, "
        "що охоплює HTML, CSS, JavaScript, React, Node.js та бази даних. "
        "Перевагами є структурований навчальний план та безкоштовність. "
        "Однак завдання є ізольованими — студент виконує окремі вправи "
        "у вбудованому редакторі, не працюючи з реальним проєктом. "
        "Відсутня інтеграція між frontend та backend частинами, немає "
        "досвіду роботи з системою контролю версій та командною розробкою."
    ))

    add_paragraph(doc, (
        "The Odin Project — безкоштовна платформа з проєктно-орієнтованим "
        "підходом. Студенти створюють реальні проєкти, використовуючи "
        "зовнішні ресурси та документацію. Перевагою є наближеність до "
        "реальної розробки. Обмеження: відсутність готового backend для "
        "frontend-студентів, немає мобільної розробки, відсутня система "
        "спринтів та поступового ускладнення."
    ))

    add_paragraph(doc, (
        "Codecademy — інтерактивна платформа з покроковими уроками. "
        "Забезпечує швидкий старт та миттєвий зворотний зв'язок. "
        "Обмеження: завдання виконуються у браузерному середовищі, "
        "що не відповідає реальним умовам розробки. Відсутній досвід "
        "роботи з повноцінним проєктом, системою збірки та залежностями."
    ))

    add_paragraph(doc, (
        "Udemy, Coursera — платформи з відеокурсами. Надають теоретичну "
        "базу та демонстрацію процесу розробки. Обмеження: пасивне "
        "навчання (перегляд відео), відсутність інтерактивності, "
        "проєкти курсу часто застарівають, немає адаптації під рівень "
        "студента."
    ))

    add_paragraph(doc, (
        "Exercism — платформа з менторством та code review. Забезпечує "
        "зворотний зв'язок від досвідчених розробників. Обмеження: "
        "фокус на алгоритмічних задачах, а не на проєктній розробці. "
        "Відсутній досвід роботи з багатокомпонентними системами."
    ))

    add_paragraph(doc, (
        "Порівняльний аналіз існуючих платформ наведено у таблиці 1.1."
    ))

    # Table 1.1
    table = doc.add_table(rows=7, cols=5)
    table.style = "Table Grid"
    headers = ["Платформа", "Реальний проєкт", "Full-stack", "Спринти", "AI-ментор"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ["freeCodeCamp", "Ні", "Частково", "Ні", "Ні"],
        ["The Odin Project", "Так", "Частково", "Ні", "Ні"],
        ["Codecademy", "Ні", "Частково", "Ні", "Так"],
        ["Udemy/Coursera", "Частково", "Так", "Ні", "Ні"],
        ["Exercism", "Ні", "Ні", "Ні", "Так"],
        ["Наша платформа", "Так", "Так", "Так", "Так"],
    ]
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_data

    add_image_caption(doc, "Таблиця 1.1 — Порівняння навчальних платформ")

    add_paragraph(doc, (
        "Як видно з таблиці, жодна з існуючих платформ не забезпечує "
        "повного набору характеристик, необхідних для ефективного навчання "
        "full-stack розробці: роботу з реальним багатокомпонентним проєктом, "
        "спринтову модель з поступовим ускладненням та AI-менторство."
    ))

    # 1.3
    add_subsection_title(doc, "1.3. Monorepo-архітектура: концепція та переваги")

    add_paragraph(doc, (
        "Monorepo (від англ. monolithic repository) — це стратегія організації "
        "коду, при якій кілька логічно незалежних проєктів зберігаються в "
        "одному репозиторії системи контролю версій. На відміну від polyrepo-підходу, "
        "де кожен проєкт має окремий репозиторій, monorepo забезпечує єдину "
        "точку входу для всієї кодової бази організації або продукту."
    ))

    add_paragraph(doc, (
        "Серед компаній, що успішно використовують monorepo-підхід: Google "
        "(понад 2 мільярди рядків коду в одному репозиторії), Meta (React, "
        "React Native та інші проєкти), Microsoft (Rush, TypeScript), "
        "Vercel (Next.js, Turborepo). Це свідчить про зрілість та "
        "масштабованість підходу."
    ))

    add_paragraph(doc, "Основні переваги monorepo-архітектури:")
    advantages = [
        "Спільний код (shared packages) — можливість виділити типи, утиліти, "
        "API-клієнти та бізнес-логіку у спільні пакети, що використовуються "
        "кількома додатками одночасно;",
        "Атомарні зміни — можливість змінити API-контракт та всіх його "
        "споживачів в одному коміті, що гарантує консистентність;",
        "Єдина система збірки — один інструмент (Turborepo, Nx, Lerna) "
        "керує збіркою всіх пакетів з урахуванням залежностей та кешування;",
        "Спрощена координація — всі розробники бачать повну картину системи, "
        "що полегшує розуміння архітектури та взаємозв'язків;",
        "Єдині стандарти — спільна конфігурація лінтерів, форматерів та "
        "TypeScript забезпечує однорідність коду.",
    ]
    for adv in advantages:
        add_list_item(doc, adv)

    add_paragraph(doc, (
        "У контексті навчальної платформи monorepo-архітектура має додаткові "
        "переваги. Студент, клонуючи один репозиторій, отримує повну систему: "
        "backend, frontend, mobile та спільні пакети. Він може запустити "
        "будь-яку комбінацію компонентів, побачити як вони взаємодіють, "
        "та зосередитись на розробці обраної частини, маючи решту як "
        "працюючу інфраструктуру."
    ))

    add_paragraph(doc, (
        "Turborepo — інструмент для управління monorepo від компанії Vercel, "
        "який забезпечує: інкрементальну збірку (перебудовує лише змінені "
        "пакети), кешування результатів збірки, паралельне виконання задач "
        "та визначення порядку збірки на основі графу залежностей. Turborepo "
        "обрано для даної платформи завдяки простоті конфігурації, швидкості "
        "роботи та інтеграції з npm workspaces."
    ))

    # 1.4
    add_subsection_title(doc, "1.4. Обґрунтування вибору технологій та інструментів")

    add_paragraph(doc, (
        "Вибір технологічного стеку для навчальної платформи обумовлений "
        "двома критеріями: відповідність сучасним індустріальним стандартам "
        "та доступність для студентів. Технології мають бути актуальними "
        "на ринку праці, добре документованими та мати активну спільноту."
    ))

    add_paragraph(doc, "Серверна частина (Backend):", bold=True)
    add_paragraph(doc, (
        "Express.js — мінімалістичний веб-фреймворк для Node.js, який є "
        "найпопулярнішим рішенням для створення REST API. Обрано завдяки "
        "простоті, гнучкості та великій кількості навчальних матеріалів. "
        "Sequelize — ORM для роботи з реляційними базами даних, що "
        "підтримує SQLite, PostgreSQL, MySQL. SQLite обрано як базу даних "
        "для спрощення розгортання — не потребує окремого сервера."
    ))

    add_paragraph(doc, "Веб-клієнт (Frontend):", bold=True)
    add_paragraph(doc, (
        "React 19 — бібліотека для побудови інтерфейсів користувача, "
        "яка домінує на ринку веб-розробки. Vite — сучасний інструмент "
        "збірки, що забезпечує миттєвий запуск dev-сервера та швидку "
        "збірку. React Router — стандартне рішення для маршрутизації "
        "у React-додатках. Redux Toolkit — бібліотека для управління "
        "станом додатку."
    ))

    add_paragraph(doc, "Мобільний клієнт (Mobile):", bold=True)
    add_paragraph(doc, (
        "React Native 0.80 — фреймворк для створення нативних мобільних "
        "додатків з використанням React. Обрано завдяки спільній кодовій "
        "базі з веб-клієнтом (React, Redux, TypeScript), що демонструє "
        "переваги monorepo-підходу. React Navigation — стандартне рішення "
        "для навігації у React Native додатках."
    ))

    add_paragraph(doc, "Спільні технології:", bold=True)
    add_paragraph(doc, (
        "TypeScript 5.x — надмножина JavaScript з статичною типізацією, "
        "що забезпечує безпеку типів та покращує досвід розробки. "
        "Використовується у всіх компонентах системи. npm workspaces — "
        "вбудований механізм Node.js для управління залежностями у "
        "monorepo. Turborepo — оркестратор збірки для monorepo."
    ))

    add_paragraph(doc, "Інструменти управління станом:", bold=True)
    add_paragraph(doc, (
        "Redux Toolkit — обрано як стандарт індустрії для управління "
        "глобальним станом. Бібліотека використовується як спільний "
        "пакет (@external-lab-monorepo/store), що демонструє "
        "перевикористання бізнес-логіки між web та mobile додатками."
    ))

    # 1.5
    add_subsection_title(doc, "1.5. Висновки за розділом")

    add_paragraph(doc, (
        "У першому розділі проведено аналіз предметної області та "
        "обґрунтовано актуальність розроблення навчальної платформи. "
        "Основні результати:"
    ))
    conclusions = [
        "Виявлено проблему фрагментарності навчання full-stack розробці "
        "та відсутності досвіду роботи з реальними багатокомпонентними системами.",
        "Проаналізовано існуючі навчальні платформи (freeCodeCamp, The Odin Project, "
        "Codecademy, Exercism) та виявлено їх обмеження: ізольованість завдань, "
        "відсутність повноцінного full-stack досвіду та спринтової моделі.",
        "Обґрунтовано вибір monorepo-архітектури як основи навчальної платформи, "
        "що забезпечує єдину точку входу, спільні пакети та цілісне розуміння системи.",
        "Обґрунтовано вибір технологічного стеку (TypeScript, React, React Native, "
        "Express, Turborepo), що відповідає сучасним індустріальним стандартам.",
    ]
    for i, c in enumerate(conclusions, 1):
        add_numbered_item(doc, i, c)


def write_chapter2(doc):
    """Write РОЗДІЛ 2 - Проєктування навчальної платформи."""
    add_page_break(doc)
    add_chapter_title(doc, "РОЗДІЛ 2. ПРОЄКТУВАННЯ НАВЧАЛЬНОЇ ПЛАТФОРМИ")

    # 2.1
    add_subsection_title(doc, "2.1. Концепція платформи та навчальна модель")

    add_paragraph(doc, (
        "Навчальна платформа побудована на концепції «навчання через реальний "
        "проєкт». На відміну від традиційних курсів, де студент виконує "
        "ізольовані вправи, дана платформа надає повноцінну працюючу систему "
        "(task manager), в якій студент розробляє одну з частин, маючи решту "
        "як готову інфраструктуру."
    ))

    add_paragraph(doc, (
        "Платформа підтримує три напрямки навчання: frontend-розробка "
        "(веб-додаток на React), backend-розробка (серверний додаток на "
        "Express) та mobile-розробка (мобільний додаток на React Native). "
        "Студент обирає напрямок та отримує:"
    ))
    items = [
        "готові працюючі компоненти системи для інших напрямків (наприклад, "
        "frontend-студент отримує готовий backend);",
        "набір навчальних завдань (тікетів), організованих у спринти з "
        "поступовим ускладненням;",
        "референсну реалізацію обраного напрямку для самоперевірки;",
        "AI-асистент, який виступає в ролі ментора та допомагає з питаннями.",
    ]
    for item in items:
        add_list_item(doc, item)

    add_paragraph(doc, (
        "Діаграму сценаріїв використання платформи наведено на рис. 2.1."
    ))
    if os.path.exists("diagrams/output/use-case.png"):
        add_image(doc, "diagrams/output/use-case.png", width=Cm(12))
    add_image_caption(doc, "Рис. 2.1 — Сценарії використання навчальної платформи")

    add_paragraph(doc, (
        "Навчальна модель базується на спринтовій методології. Кожен спринт "
        "додає новий рівень складності та нові концепції:"
    ))
    sprints = [
        "Спринт 1 — Авторизація: реєстрація, вхід, робота з JWT-токенами, "
        "захищені маршрути;",
        "Спринт 2 — Читання даних: відображення списків, профіль користувача, "
        "обробка помилок, стани завантаження;",
        "Спринт 3 — Повний CRUD: створення, редагування, видалення задач, "
        "робота з файлами, пагінація;",
        "Спринт 4 — Розширення функціоналу: пріоритети задач, оновлення "
        "дизайну, додаткові фільтри.",
    ]
    for sprint in sprints:
        add_list_item(doc, sprint)

    # 2.2
    add_subsection_title(doc, "2.2. Сценарії використання платформи")

    add_paragraph(doc, (
        "Для демонстрації практичної цінності платформи розглянемо основні "
        "сценарії її використання."
    ))

    add_paragraph(doc, (
        "Сценарій 1. Студент вивчає Frontend-розробку."
    ), bold=True)
    add_paragraph(doc, (
        "Студент клонує monorepo-репозиторій та запускає серверний додаток "
        "командою у терміналі. Backend стартує на порту 4000 та надає "
        "повноцінний REST API. Студент відкриває документацію з тікетами "
        "для Frontend Sprint 1, створює свій окремий React-проєкт та "
        "починає реалізовувати задачі: верстку сторінки реєстрації, "
        "валідацію форм, інтеграцію з API авторизації. Його додаток "
        "працює з готовим backend, що дозволяє одразу бачити результат. "
        "Якщо студент стикається з труднощами, він може звернутися до "
        "AI-асистента або переглянути референсну реалізацію frontend-додатку "
        "у monorepo."
    ))

    add_paragraph(doc, (
        "Сценарій 2. Студент вивчає Backend-розробку."
    ), bold=True)
    add_paragraph(doc, (
        "Студент запускає готовий frontend-додаток з monorepo. Веб-клієнт "
        "стартує та очікує API на визначених ендпоінтах. Студент відкриває "
        "тікети для Backend Sprint 1, створює свій Express-проєкт та "
        "реалізує API авторизації згідно з контрактами. Готовий frontend "
        "автоматично працює з його backend, що дає миттєвий візуальний "
        "зворотний зв'язок — студент бачить як його API використовується "
        "реальним клієнтським додатком. Додатково студент може запустити "
        "мобільний додаток і переконатися, що його API працює і з ним."
    ))

    add_paragraph(doc, (
        "Сценарій 3. Студент вивчає Mobile-розробку."
    ), bold=True)
    add_paragraph(doc, (
        "Студент запускає backend з monorepo та відкриває тікети для "
        "Mobile Sprint 1. Він створює свій React Native проєкт та "
        "реалізує мобільні екрани, працюючи з тим самим API, що й "
        "веб-клієнт. Готовий frontend слугує візуальним референсом — "
        "студент бачить як має виглядати та працювати додаток на вебі "
        "і реалізує аналогічний функціонал для мобільної платформи."
    ))

    add_paragraph(doc, (
        "Сценарій 4. Взаємодія з AI-асистентом."
    ), bold=True)
    add_paragraph(doc, (
        "На будь-якому етапі роботи студент може звернутися до AI-асистента. "
        "Асистент має контекст проєкту: знає структуру monorepo, технологічний "
        "стек, зміст тікетів поточного спринту та архітектурні рішення. "
        "Студент запитує: «Як реалізувати захищений маршрут у React?» — "
        "асистент пояснює концепцію, дає підказки щодо підходу, але не "
        "надає готовий код. Це імітує роботу з ментором у реальній компанії."
    ))

    # 2.3
    add_subsection_title(doc, "2.3. Архітектура Monorepo-системи")

    add_paragraph(doc, (
        "Архітектура платформи побудована на основі npm workspaces та "
        "Turborepo. Репозиторій має наступну структуру верхнього рівня:"
    ))
    add_list_item(doc, "apps/ — директорія з додатками (frontend-app, backend-app, mobile-app);")
    add_list_item(doc, "packages/ — директорія зі спільними пакетами (types, api, store, hooks, constants);")
    add_list_item(doc, "docs/ — документація та діаграми;")
    add_list_item(doc, "turbo.json — конфігурація Turborepo;")
    add_list_item(doc, "package.json — кореневий файл з визначенням workspaces.")

    add_paragraph(doc, (
        "Конфігурація npm workspaces у кореневому package.json визначає "
        "два workspace-шляхи: «apps/*» та «packages/*». Це дозволяє npm "
        "автоматично зв'язувати локальні пакети між собою без публікації "
        "у реєстр. Кожен додаток може імпортувати спільні пакети за їх "
        "іменем (наприклад, @external-lab-monorepo/types) так само, як "
        "зовнішні npm-залежності."
    ))

    add_paragraph(doc, (
        "Turborepo конфігурується через turbo.json, де визначено дві задачі: "
        "«build» (з залежністю від збірки пакетів-залежностей через «^build») "
        "та «dev» (без кешування, з прапорцем persistent для dev-серверів). "
        "Це забезпечує правильний порядок збірки: спочатку збираються "
        "спільні пакети (types → constants → api → store → hooks), потім "
        "додатки, що від них залежать."
    ))

    add_paragraph(doc, (
        "Діаграму залежностей між компонентами monorepo наведено на рис. 2.2."
    ))
    # Image placeholder
    if os.path.exists("diagrams/output/monorepo-structure.png"):
        add_image(doc, "diagrams/output/monorepo-structure.png", width=Cm(13))
    add_image_caption(doc, "Рис. 2.2 — Структура monorepo та залежності між пакетами")

    # 2.4
    add_subsection_title(doc, "2.4. Проєктування спільних пакетів")

    add_paragraph(doc, (
        "Спільні пакети (shared packages) — ключовий елемент архітектури "
        "платформи, що демонструє переваги monorepo-підходу. Кожен пакет "
        "має подвійну збірку (ESM та CJS) для сумісності з різними "
        "середовищами виконання."
    ))

    add_paragraph(doc, "@external-lab-monorepo/types — пакет типів:", bold=True)
    add_paragraph(doc, (
        "Містить TypeScript-інтерфейси для основних сутностей системи: "
        "Task (задача з полями id, title, description, done, priority, files), "
        "CommonTask (спільна задача), User (користувач з полями email, name, "
        "avatar) та TaskPriority (тип пріоритету). Цей пакет є фундаментом "
        "типобезпеки — всі додатки та інші пакети імпортують типи з нього, "
        "що гарантує консистентність даних між frontend, backend та mobile."
    ))

    add_paragraph(doc, "@external-lab-monorepo/constants — пакет констант:", bold=True)
    add_paragraph(doc, (
        "Містить спільні константи: enum SPRINTS (визначає номери спринтів "
        "від 1 до 4) та TASKS_PER_PAGE (кількість задач на сторінку для "
        "пагінації). Винесення констант у спільний пакет забезпечує єдине "
        "джерело правди для конфігурації."
    ))

    add_paragraph(doc, "@external-lab-monorepo/api — пакет API-клієнта:", bold=True)
    add_paragraph(doc, (
        "Містить Axios-інстанс з налаштованим interceptor для автоматичного "
        "додавання JWT-токена до запитів, а також функції для кожного "
        "API-ендпоінту: авторизація (signUp, signIn), задачі (getTasks, "
        "createTask, updateTask, deleteTask), користувач (getProfile, "
        "updateProfile). Пакет залежить від @external-lab-monorepo/types "
        "для типізації запитів та відповідей."
    ))

    add_paragraph(doc, "@external-lab-monorepo/store — пакет стану:", bold=True)
    add_paragraph(doc, (
        "Містить Redux Toolkit store з чотирма слайсами: authReducer "
        "(стан авторизації), tasksReducer (список задач), currentUserReducer "
        "(профіль користувача), commonTasksReducer (спільні задачі). "
        "Кожен слайс має відповідні async thunks для взаємодії з API "
        "та селектори для отримання даних. Пакет залежить від "
        "@external-lab-monorepo/api та @external-lab-monorepo/types."
    ))

    add_paragraph(doc, "@external-lab-monorepo/hooks — пакет хуків:", bold=True)
    add_paragraph(doc, (
        "Містить React-хуки, що інкапсулюють роботу зі store: useAuth "
        "(авторизація та вихід), useTasks (список задач з пагінацією), "
        "useTask (окрема задача), useCurrentUser (профіль), useCommonTasks "
        "(спільні задачі). Хуки надають зручний інтерфейс для компонентів, "
        "приховуючи деталі роботи з Redux."
    ))

    # 2.5
    add_subsection_title(doc, "2.5. Проєктування бази даних")

    add_paragraph(doc, (
        "База даних платформи побудована на SQLite з використанням ORM "
        "Sequelize. Модель даних включає три основні сутності: User "
        "(користувач), Task (задача) та File (файл-вкладення)."
    ))

    add_paragraph(doc, (
        "Сутність User містить поля: id (первинний ключ, автоінкремент), "
        "email (унікальний, обов'язковий), name (обов'язковий), password "
        "(обов'язковий, зберігається у хешованому вигляді через bcrypt), "
        "avatar (необов'язковий, шлях до файлу зображення)."
    ))

    add_paragraph(doc, (
        "Сутність Task містить поля: id (первинний ключ), title (обов'язковий), "
        "description (необов'язковий), done (булевий, за замовчуванням false), "
        "priority (enum: 'high' або 'low', необов'язковий), userId (зовнішній "
        "ключ на User)."
    ))

    add_paragraph(doc, (
        "Сутність File містить поля: id (первинний ключ), image (шлях до "
        "файлу, обов'язковий), taskId (зовнішній ключ на Task)."
    ))

    add_paragraph(doc, (
        "Зв'язки між сутностями: User має багато Task (один-до-багатьох, "
        "каскадне видалення), Task має багато File (один-до-багатьох, "
        "каскадне видалення). Також існує сутність CommonTask з полями "
        "id та title, яка не пов'язана з користувачами і містить "
        "загальнодоступні задачі."
    ))

    add_paragraph(doc, "ER-діаграму бази даних наведено на рис. 2.3.")
    if os.path.exists("diagrams/output/er-model.png"):
        add_image(doc, "diagrams/output/er-model.png", width=Cm(12))
    add_image_caption(doc, "Рис. 2.3 — ER-модель бази даних")

    # 2.6
    add_subsection_title(doc, "2.6. Проєктування REST API та спринтовий роутинг")

    add_paragraph(doc, (
        "Серверний додаток надає REST API за адресою /api. Ключовою "
        "архітектурною особливістю є спринтовий роутинг — механізм, "
        "який дозволяє одному серверу обслуговувати різні версії API "
        "залежно від спринту студента."
    ))

    add_paragraph(doc, (
        "Клієнт передає номер спринту через HTTP-заголовок «sprint». "
        "Серверний роутер аналізує цей заголовок та направляє запит "
        "до відповідного набору контролерів. Кожен спринт має свою "
        "директорію з controllers, routes та services, що реалізують "
        "функціонал відповідного рівня складності:"
    ))
    add_list_item(doc, "Sprint 1: POST /api/auth/sign-up, POST /api/auth/sign-in;")
    add_list_item(doc, "Sprint 2: + GET /api/tasks, GET /api/profile, PUT /api/profile;")
    add_list_item(doc, "Sprint 3: + POST /api/tasks, PUT /api/tasks/:id, DELETE /api/tasks/:id, GET /api/tasks/:id, GET /api/common-tasks;")
    add_list_item(doc, "Sprint 4: + GET /api/priorities, оновлені ендпоінти з підтримкою пріоритетів.")

    add_paragraph(doc, (
        "Такий підхід дозволяє backend-студенту поступово реалізовувати "
        "API, а frontend/mobile-студенту — працювати з API відповідного "
        "рівня складності, не стикаючись з ендпоінтами, які ще не "
        "вивчались."
    ))

    add_paragraph(doc, (
        "Серверний додаток також включає middleware-шар: authMiddleware "
        "(перевірка JWT-токена), adminRoleCheckMiddleware (перевірка прав "
        "адміністратора), errorHandlingMiddleware (централізована обробка "
        "помилок), validators (валідація вхідних даних через express-validator)."
    ))

    # 2.7
    add_subsection_title(doc, "2.7. Проєктування web-додатку")

    add_paragraph(doc, (
        "Web-додаток побудований на React з використанням Vite як "
        "інструменту збірки. Архітектура додатку базується на патерні "
        "Container/Presentational: кожна сторінка розділена на Container "
        "(логіка, хуки, стан) та Page (чистий UI через пропси)."
    ))

    add_paragraph(doc, "Структура web-додатку включає:")
    add_list_item(doc, "components/ — перевикористовувані UI-компоненти (CustomButton, TextInputWithHint, UserAvatar, TaskListItem тощо);")
    add_list_item(doc, "pages/ — сторінки додатку, розділені на auth (sign-in, sign-up) та app (tasks, profile, add-task, edit-task, task-details, common-tasks);")
    add_list_item(doc, "router/ — маршрутизація з GuestRoute (для неавторизованих) та ProtectedRoute (для авторизованих);")
    add_list_item(doc, "layouts/ — макети сторінок (AuthLayout для авторизації, AppLayout з навігацією);")
    add_list_item(doc, "contexts/ — React-контексти (DevMenuContext для перемикання спринтів).")

    add_paragraph(doc, (
        "DevMenu — спеціальний компонент для розробки, який дозволяє "
        "перемикати поточний спринт. При зміні спринту змінюється "
        "HTTP-заголовок, що надсилається до backend, і відповідно "
        "змінюється доступний функціонал API."
    ))

    add_paragraph(doc, "Діаграму компонентів web-додатку наведено на рис. 2.4.")
    if os.path.exists("diagrams/output/component-tree.png"):
        add_image(doc, "diagrams/output/component-tree.png", width=Cm(15))
    add_image_caption(doc, "Рис. 2.4 — Дерево компонентів web-додатку")

    # 2.8
    add_subsection_title(doc, "2.8. Проєктування mobile-додатку")

    add_paragraph(doc, (
        "Мобільний додаток побудований на React Native з використанням "
        "React Navigation для навігації. Архітектура аналогічна web-додатку "
        "(Container/Presentational), що демонструє перевикористання "
        "архітектурних патернів між платформами."
    ))

    add_paragraph(doc, "Структура mobile-додатку включає:")
    add_list_item(doc, "components/ — нативні UI-компоненти;")
    add_list_item(doc, "screens/ — екрани додатку (аналоги сторінок web-версії);")
    add_list_item(doc, "navigation/ — конфігурація навігації (Bottom Tabs + Stack Navigator);")
    add_list_item(doc, "hooks/ — платформо-специфічні хуки;")
    add_list_item(doc, "constants/ — кольори, шрифти, розміри.")

    add_paragraph(doc, (
        "Мобільний додаток використовує ті самі спільні пакети "
        "(@external-lab-monorepo/hooks, store, api, types, constants), "
        "що й web-додаток. Це демонструє головну перевагу monorepo — "
        "бізнес-логіка (авторизація, робота з задачами, управління станом) "
        "написана один раз і використовується на обох платформах."
    ))

    add_paragraph(doc, "Діаграму навігації mobile-додатку наведено на рис. 2.5.")
    if os.path.exists("diagrams/output/navigation-flow.png"):
        add_image(doc, "diagrams/output/navigation-flow.png", width=Cm(13))
    add_image_caption(doc, "Рис. 2.5 — Навігаційний потік додатку")

    # 2.9
    add_subsection_title(doc, "2.9. Система навчальних завдань (тікети та спринти)")

    add_paragraph(doc, (
        "Навчальні завдання оформлені у вигляді тікетів — структурованих "
        "документів, що імітують реальні задачі в системах управління "
        "проєктами (Jira, Linear, GitHub Issues). Кожен тікет містить:"
    ))
    add_list_item(doc, "Epic — група пов'язаних задач (Authorization, Tasks, Profile, Redesign, Task Priority);")
    add_list_item(doc, "Тип — Story (функціональна задача) або Task (технічна задача);")
    add_list_item(doc, "Назва — короткий опис задачі;")
    add_list_item(doc, "Objective — мета задачі;")
    add_list_item(doc, "Details — технічні деталі (шлях, посилання на дизайн);")
    add_list_item(doc, "Scenarios — детальні сценарії поведінки (happy path, error handling, edge cases).")

    add_paragraph(doc, (
        "Тікети розподілені по спринтах для кожного напрямку навчання. "
        "Для Frontend-напрямку передбачено 27 тікетів, для Backend — 31, "
        "для Mobile — 27. Кожен набір охоплює повний цикл розробки "
        "task manager додатку від авторизації до розширеного функціоналу."
    ))

    add_paragraph(doc, (
        "Діаграму спринтової прогресії наведено на рис. 2.6."
    ))
    if os.path.exists("diagrams/output/sprint-progression.png"):
        add_image(doc, "diagrams/output/sprint-progression.png", width=Cm(4.5))
    add_image_caption(doc, "Рис. 2.6 — Спринтова прогресія навчальних завдань")

    add_paragraph(doc, (
        "Формат тікетів обрано свідомо — він відповідає реальним практикам "
        "індустрії. Студент не лише вчиться програмувати, а й набуває "
        "навичок роботи з технічною документацією, розуміння acceptance "
        "criteria та сценаріїв тестування."
    ))

    # 2.10
    add_subsection_title(doc, "2.10. Проєктування AI-асистента навчального процесу")

    add_paragraph(doc, (
        "AI-асистент — інтегрований компонент платформи, що виступає "
        "в ролі ментора для студента. Асистент побудований на основі "
        "великої мовної моделі (LLM) та має доступ до контексту проєкту."
    ))

    add_paragraph(doc, "Архітектура AI-асистента включає:")
    add_list_item(doc, "Контекстне вікно — містить інформацію про структуру monorepo, технологічний стек, поточний спринт та зміст тікетів;")
    add_list_item(doc, "Системний промпт — визначає роль асистента (ментор, не виконавець), обмеження (не давати готовий код) та стиль відповідей;")
    add_list_item(doc, "База знань — документація проєкту, архітектурні рішення, типові помилки та їх вирішення.")

    add_paragraph(doc, "Сценарії взаємодії з AI-асистентом:")
    add_list_item(doc, "Пояснення концепцій — студент запитує «Що таке JWT?» і отримує пояснення в контексті проєкту;")
    add_list_item(doc, "Підказки щодо реалізації — студент запитує «Як реалізувати захищений маршрут?» і отримує напрямок без готового коду;")
    add_list_item(doc, "Допомога з помилками — студент описує помилку і отримує пояснення причини та підказку щодо вирішення;")
    add_list_item(doc, "Навігація по проєкту — студент запитує «Де знаходиться конфігурація API?» і отримує пояснення структури.")

    add_paragraph(doc, (
        "Ключовий принцип AI-асистента — він направляє, а не виконує. "
        "Асистент пояснює підходи, дає підказки, задає навідні питання, "
        "але не генерує готовий код для копіювання. Це забезпечує "
        "активне навчання та розвиток самостійності студента."
    ))

    add_paragraph(doc, (
        "Діаграму взаємодії студента з AI-асистентом наведено на рис. 2.7."
    ))
    if os.path.exists("diagrams/output/ai-assistant-flow.png"):
        add_image(doc, "diagrams/output/ai-assistant-flow.png", width=Cm(13))
    add_image_caption(doc, "Рис. 2.7 — Схема взаємодії студента з AI-асистентом")

    # 2.11
    add_subsection_title(doc, "2.11. Висновки за розділом")

    add_paragraph(doc, "У другому розділі виконано проєктування навчальної платформи:")
    conclusions = [
        "Сформульовано концепцію платформи та навчальну модель на основі спринтової методології з поступовим ускладненням.",
        "Описано сценарії використання для трьох напрямків навчання (frontend, backend, mobile) та взаємодії з AI-асистентом.",
        "Спроєктовано архітектуру monorepo-системи з використанням Turborepo та npm workspaces.",
        "Спроєктовано п'ять спільних пакетів (types, constants, api, store, hooks), що забезпечують перевикористання коду.",
        "Спроєктовано базу даних з трьома основними сутностями та їх зв'язками.",
        "Спроєктовано REST API зі спринтовим роутингом через HTTP-заголовок.",
        "Спроєктовано web та mobile додатки з архітектурою Container/Presentational.",
        "Розроблено систему навчальних завдань у форматі тікетів.",
        "Спроєктовано AI-асистент навчального процесу на основі LLM.",
    ]
    for i, c in enumerate(conclusions, 1):
        add_numbered_item(doc, i, c)


def write_chapter3(doc):
    """Write РОЗДІЛ 3 - Реалізація та тестування."""
    add_page_break(doc)
    add_chapter_title(doc, "РОЗДІЛ 3. РЕАЛІЗАЦІЯ ТА ТЕСТУВАННЯ")

    # 3.1
    add_subsection_title(doc, "3.1. Реалізація серверної частини")

    add_paragraph(doc, (
        "Серверний додаток реалізовано на Express.js 5 з використанням "
        "TypeScript. Точка входу (index.ts) ініціалізує Express-додаток, "
        "підключає middleware (CORS, JSON-парсер, обробку файлів, "
        "статичні файли) та запускає сервер на порту 4000. Після запуску "
        "виконується підключення до бази даних SQLite через Sequelize "
        "з механізмом повторних спроб (до 5 спроб з інтервалом 1 секунда)."
    ))

    add_paragraph(doc, (
        "Моделі даних визначено через Sequelize з використанням TypeScript-"
        "інтерфейсів для типобезпеки. Модель User визначає поля email "
        "(unique), name, password та avatar. Модель Task — title, description, "
        "done, priority з зовнішнім ключем userId. Модель File — image "
        "з зовнішнім ключем taskId. Зв'язки налаштовано з каскадним "
        "видаленням."
    ))

    add_paragraph(doc, (
        "Спринтовий роутинг реалізовано через центральний роутер, який "
        "зчитує заголовок «sprint» з HTTP-запиту та делегує обробку "
        "відповідному набору маршрутів. Кожен спринт має власну "
        "директорію з трирівневою архітектурою: routes (визначення "
        "ендпоінтів), controllers (обробка запитів), services (бізнес-логіка)."
    ))

    add_paragraph(doc, (
        "Автентифікація реалізована через JWT (jsonwebtoken). При реєстрації "
        "пароль хешується через bcrypt, при вході — порівнюється з хешем. "
        "Успішна авторизація повертає JWT-токен, який клієнт передає "
        "у заголовку Authorization для доступу до захищених ендпоінтів. "
        "authMiddleware перевіряє валідність токена та додає userId до "
        "об'єкта запиту."
    ))

    add_paragraph(doc, (
        "Завантаження файлів реалізовано через express-fileupload. Файли "
        "зберігаються у директорії static/ з унікальними іменами (UUID). "
        "Валідація вхідних даних виконується через express-validator "
        "з кастомними правилами для кожного ендпоінту."
    ))

    add_paragraph(doc, (
        "Архітектуру серверного додатку наведено на рис. 3.1."
    ))
    if os.path.exists("diagrams/output/backend-layers.png"):
        add_image(doc, "diagrams/output/backend-layers.png", width=Cm(12))
    add_image_caption(doc, "Рис. 3.1 — Архітектура серверного додатку")

    # 3.2
    add_subsection_title(doc, "3.2. Реалізація спільних пакетів")

    add_paragraph(doc, (
        "Кожен спільний пакет має подвійну збірку: ESM (для web-додатку "
        "з Vite) та CJS (для backend та React Native). Збірка виконується "
        "через TypeScript compiler (tsc) з двома конфігураціями: "
        "tsconfig.esm.json та tsconfig.cjs.json. Результати збірки "
        "розміщуються у dist/esm/ та dist/cjs/ відповідно."
    ))

    add_paragraph(doc, (
        "Пакет types експортує інтерфейси Task, CommonTask, User та тип "
        "TaskPriority. Ці типи використовуються у всіх інших пакетах "
        "та додатках для забезпечення типобезпеки на рівні всієї системи."
    ))

    add_paragraph(doc, (
        "Пакет api реалізує Axios-інстанс з базовою URL-адресою та "
        "interceptor, який автоматично додає JWT-токен з localStorage "
        "(для web) або AsyncStorage (для mobile) до кожного запиту. "
        "Функції API типізовані: кожна приймає та повертає типи з "
        "пакету types."
    ))

    add_paragraph(doc, (
        "Пакет store реалізує Redux Toolkit store з чотирма слайсами. "
        "Кожен слайс використовує createAsyncThunk для асинхронних "
        "операцій та createSlice для визначення reducers. Селектори "
        "реалізовані через reselect для мемоізації та оптимізації "
        "перерендерів."
    ))

    add_paragraph(doc, (
        "Пакет hooks надає React-хуки, що інкапсулюють dispatch та "
        "useSelector. Наприклад, useAuth повертає функції signIn, "
        "signUp, logout та стан isAuthenticated, isLoading, error. "
        "Це спрощує використання store у компонентах та забезпечує "
        "єдиний інтерфейс для web та mobile."
    ))

    # 3.3
    add_subsection_title(doc, "3.3. Реалізація web-додатку")

    add_paragraph(doc, (
        "Web-додаток реалізовано на React 19 з Vite як інструментом "
        "збірки. Маршрутизація побудована на React Router DOM 7 з "
        "двома типами маршрутів: GuestRoute (доступний лише "
        "неавторизованим, редирект на / при наявності токена) та "
        "ProtectedRoute (доступний лише авторизованим, редирект на "
        "/sign-in при відсутності токена)."
    ))

    add_paragraph(doc, (
        "Кожна сторінка реалізована за патерном Container/Presentational. "
        "Container-компонент використовує хуки зі спільного пакету "
        "(@external-lab-monorepo/hooks), управляє станом та передає "
        "дані через пропси у Page-компонент, який відповідає лише "
        "за відображення UI."
    ))

    add_paragraph(doc, (
        "DevMenu — компонент для розробки, реалізований через React Context. "
        "Дозволяє перемикати поточний спринт через випадаючий список. "
        "При зміні спринту оновлюється заголовок HTTP-запитів, що "
        "змінює поведінку backend. Це дозволяє студенту тестувати "
        "свій frontend з різними рівнями API."
    ))

    add_paragraph(doc, (
        "Компоненти UI реалізовані з використанням CSS-модулів. "
        "Кожен компонент має окремий CSS-файл з відповідними стилями. "
        "Реалізовано компоненти: CustomButton (кнопка з варіантами), "
        "TextInputWithHint (поле вводу з підказкою та помилкою), "
        "UserAvatar (аватар з можливістю завантаження), UniversalLoading "
        "(індикатор завантаження), UniversalError (відображення помилок), "
        "TaskListItem (елемент списку задач), TaskPriority (відображення "
        "пріоритету), TaskStatus (статус задачі)."
    ))

    add_paragraph(doc, "Діаграму потоків даних у web-додатку наведено на рис. 3.2.")
    if os.path.exists("diagrams/output/data-flow.png"):
        add_image(doc, "diagrams/output/data-flow.png", width=Cm(6))
    add_image_caption(doc, "Рис. 3.2 — Потік даних у додатку")

    # 3.4
    add_subsection_title(doc, "3.4. Реалізація mobile-додатку")

    add_paragraph(doc, (
        "Мобільний додаток реалізовано на React Native 0.80 з "
        "використанням React Navigation 5 для навігації. Навігаційна "
        "структура включає Bottom Tab Navigator (вкладки: Tasks, "
        "Profile) та Stack Navigator для вкладених екранів (TaskDetails, "
        "AddTask, EditTask)."
    ))

    add_paragraph(doc, (
        "Додаток використовує ті самі спільні пакети, що й web-версія. "
        "Це означає, що бізнес-логіка (авторизація, робота з задачами, "
        "управління станом) ідентична на обох платформах. Відрізняється "
        "лише UI-шар: замість HTML-елементів використовуються нативні "
        "компоненти React Native (View, Text, TextInput, TouchableOpacity, "
        "FlatList тощо)."
    ))

    add_paragraph(doc, (
        "Стилізація реалізована через StyleSheet.create() — нативний "
        "механізм React Native для визначення стилів. Тема додатку "
        "(кольори, шрифти, відступи) визначена у constants/ та "
        "використовується через всі компоненти для забезпечення "
        "консистентного дизайну."
    ))

    add_paragraph(doc, (
        "Навігація між екранами авторизації та основним додатком "
        "контролюється через стан авторизації у Redux store. При "
        "наявності токена відображається основний Stack з вкладками, "
        "при відсутності — екрани Sign In / Sign Up."
    ))

    # 3.5
    add_subsection_title(doc, "3.5. Реалізація AI-асистента")

    add_paragraph(doc, (
        "AI-асистент реалізовано як інтеграцію з великою мовною моделлю "
        "(LLM), яка отримує контекст навчальної платформи через системний "
        "промпт. Контекст включає:"
    ))
    add_list_item(doc, "структуру monorepo та призначення кожного пакету;")
    add_list_item(doc, "технологічний стек та версії бібліотек;")
    add_list_item(doc, "зміст тікетів поточного спринту;")
    add_list_item(doc, "архітектурні рішення та патерни, використані у проєкті;")
    add_list_item(doc, "типові помилки студентів та способи їх вирішення.")

    add_paragraph(doc, (
        "Системний промпт визначає поведінку асистента: він виступає "
        "в ролі досвідченого ментора, який пояснює концепції, дає "
        "підказки та направляє студента, але не генерує готовий код "
        "для копіювання. Асистент адаптує рівень пояснень залежно "
        "від спринту студента — для Sprint 1 пояснення більш детальні, "
        "для Sprint 4 — більш лаконічні."
    ))

    add_paragraph(doc, (
        "Технічно AI-асистент може бути реалізований через різні LLM-"
        "провайдери (OpenAI API, Anthropic Claude, локальні моделі через "
        "Ollama). Платформа надає підготовлений контекст у форматі "
        "markdown-файлів, які можуть бути завантажені у будь-який "
        "LLM-інтерфейс (ChatGPT, Claude, IDE-інтеграції типу Cursor "
        "або GitHub Copilot)."
    ))

    # 3.6
    add_subsection_title(doc, "3.6. Інтеграція компонентів та демонстрація роботи платформи")

    add_paragraph(doc, (
        "Інтеграція компонентів забезпечується через npm workspaces та "
        "Turborepo. Команда «npm run dev» у кореневій директорії запускає "
        "всі три додатки паралельно: backend на порту 4000, frontend "
        "на порту 5173 (Vite dev server), mobile через Metro bundler."
    ))

    add_paragraph(doc, (
        "Для демонстрації роботи платформи розглянемо типовий сценарій "
        "frontend-студента на Sprint 2. Студент запускає backend командою "
        "у директорії apps/backend-app. Сервер стартує, підключається "
        "до SQLite та виводить повідомлення про готовність. Студент "
        "відкриває тікет «My Tasks Page implementation», читає сценарії "
        "та починає реалізацію у своєму проєкті. Його React-додаток "
        "надсилає GET-запит на localhost:4000/api/tasks з заголовком "
        "sprint: 2 та отримує список задач у JSON-форматі."
    ))

    add_paragraph(doc, (
        "Якщо студент хоче побачити як має працювати готовий додаток, "
        "він запускає frontend з monorepo та бачить повноцінний "
        "працюючий web-клієнт. Аналогічно для mobile — запуск "
        "мобільного додатку демонструє очікувану поведінку на "
        "мобільній платформі."
    ))

    # 3.7
    add_subsection_title(doc, "3.7. Тестування")

    add_paragraph(doc, (
        "Тестування платформи проведено на кількох рівнях для забезпечення "
        "якості та надійності всіх компонентів."
    ))

    add_paragraph(doc, "Функціональне тестування backend:", bold=True)
    add_paragraph(doc, (
        "Перевірено коректність роботи всіх API-ендпоінтів для кожного "
        "спринту. Тестування виконано через Postman з колекціями запитів "
        "для кожного спринту. Перевірено: реєстрацію та авторизацію, "
        "CRUD-операції з задачами, завантаження файлів, пагінацію, "
        "обробку помилок та валідацію."
    ))

    add_paragraph(doc, "Інтеграційне тестування:", bold=True)
    add_paragraph(doc, (
        "Перевірено взаємодію між frontend/mobile та backend. Тестовано "
        "сценарії: авторизація → отримання токена → запит захищених "
        "ресурсів → відображення даних. Перевірено коректність роботи "
        "спринтового роутингу при зміні заголовка sprint."
    ))

    add_paragraph(doc, "Тестування спільних пакетів:", bold=True)
    add_paragraph(doc, (
        "Перевірено коректність збірки пакетів у ESM та CJS форматах. "
        "Перевірено імпорт пакетів у всіх трьох додатках. Перевірено "
        "типобезпеку — TypeScript compiler виявляє помилки типів "
        "на етапі збірки."
    ))

    add_paragraph(doc, "Тестування навчального процесу:", bold=True)
    add_paragraph(doc, (
        "Проведено пілотне тестування з групою студентів. Перевірено "
        "зрозумілість тікетів, достатність документації, коректність "
        "роботи AI-асистента та відповідність складності спринтів "
        "рівню підготовки студентів."
    ))

    # 3.8
    add_subsection_title(doc, "3.8. Висновки за розділом")

    add_paragraph(doc, "У третьому розділі виконано реалізацію та тестування платформи:")
    conclusions = [
        "Реалізовано серверний додаток на Express.js з JWT-авторизацією, Sequelize ORM та спринтовим роутингом.",
        "Реалізовано п'ять спільних пакетів з подвійною збіркою (ESM/CJS) для сумісності з web та mobile.",
        "Реалізовано web-додаток на React з Vite, React Router та патерном Container/Presentational.",
        "Реалізовано mobile-додаток на React Native з React Navigation та спільними пакетами.",
        "Реалізовано AI-асистент на основі LLM з контекстом проєкту.",
        "Проведено функціональне, інтеграційне тестування та пілотне тестування навчального процесу.",
    ]
    for i, c in enumerate(conclusions, 1):
        add_numbered_item(doc, i, c)


def write_conclusions(doc):
    """Write ЗАГАЛЬНІ ВИСНОВКИ."""
    add_page_break(doc)
    add_chapter_title(doc, "ЗАГАЛЬНІ ВИСНОВКИ")

    add_paragraph(doc, (
        "У бакалаврській кваліфікаційній роботі розроблено навчальну "
        "платформу для набуття практичних навичок full-stack розробки "
        "на основі monorepo-архітектури. В процесі виконання роботи "
        "отримано наступні результати:"
    ))

    conclusions = [
        "Проведено аналіз існуючих навчальних платформ для розробників "
        "(freeCodeCamp, The Odin Project, Codecademy, Exercism) та виявлено "
        "їх ключові обмеження: ізольованість завдань, відсутність досвіду "
        "роботи з багатокомпонентними системами, відсутність спринтової "
        "моделі навчання.",

        "Обґрунтовано вибір monorepo-архітектури як основи навчальної "
        "платформи. Monorepo забезпечує єдину точку входу для всієї "
        "системи, спільні пакети для перевикористання коду та можливість "
        "студенту бачити повну картину проєкту.",

        "Спроєктовано та реалізовано навчальну платформу, що включає "
        "серверний додаток (Express.js, Sequelize, SQLite), веб-клієнт "
        "(React, Vite, Redux), мобільний клієнт (React Native) та п'ять "
        "спільних пакетів (types, constants, api, store, hooks), об'єднаних "
        "у monorepo з використанням Turborepo.",

        "Розроблено спринтову систему навчальних завдань (4 спринти, "
        "понад 80 тікетів для трьох напрямків), що забезпечує поступове "
        "ускладнення від авторизації до повного CRUD з файлами та "
        "пріоритетами.",

        "Реалізовано унікальний механізм спринтового роутингу на backend, "
        "який дозволяє одному серверу обслуговувати різні рівні API "
        "залежно від спринту студента.",

        "Інтегровано AI-асистент навчального процесу на основі LLM, "
        "який виступає в ролі ментора: пояснює концепції, дає підказки "
        "та направляє студента без надання готового коду.",

        "Проведено тестування платформи на функціональному та "
        "інтеграційному рівнях, а також пілотне тестування навчального "
        "процесу, що підтвердило ефективність обраного підходу.",
    ]
    for i, c in enumerate(conclusions, 1):
        add_numbered_item(doc, i, c)

    add_paragraph(doc, (
        "Розроблена платформа може бути використана у вищих навчальних "
        "закладах для курсів веб- та мобільної розробки, на курсах "
        "підвищення кваліфікації та для самостійного навчання. "
        "Подальший розвиток платформи може включати: додавання нових "
        "спринтів з більш складним функціоналом, інтеграцію системи "
        "автоматичної перевірки коду, розширення AI-асистента "
        "можливостями code review та адаптацію під інші технологічні "
        "стеки."
    ))


def write_references(doc):
    """Write СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ."""
    add_page_break(doc)
    add_chapter_title(doc, "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ")

    references = [
        "Potvin R., Levenberg J. Why Google Stores Billions of Lines of Code in a Single Repository. Communications of the ACM. 2016. Vol. 59, No. 7. P. 78–87.",
        "Brousse N. The issue of monorepo and polyrepo in large enterprises. Companion Proceedings of the 2019 ACM SIGPLAN International Conference on Systems, Programming, Languages, and Applications. 2019. P. 1–4.",
        "Freeman A. Pro React 16. Apress, 2019. 745 p.",
        "Eisenman B. Learning React Native: Building Native Mobile Apps with JavaScript. O'Reilly Media, 2017. 242 p.",
        "Brown E. Web Development with Node and Express. O'Reilly Media, 2019. 330 p.",
        "Vanderkam D. Effective TypeScript: 62 Specific Ways to Improve Your TypeScript. O'Reilly Media, 2019. 264 p.",
        "Abramov D. Redux Documentation. URL: https://redux.js.org/ (дата звернення: 15.05.2026).",
        "Turborepo Documentation. URL: https://turbo.build/repo/docs (дата звернення: 10.05.2026).",
        "React Documentation. URL: https://react.dev/ (дата звернення: 12.05.2026).",
        "React Native Documentation. URL: https://reactnative.dev/ (дата звернення: 12.05.2026).",
        "Express.js Documentation. URL: https://expressjs.com/ (дата звернення: 10.05.2026).",
        "Sequelize Documentation. URL: https://sequelize.org/ (дата звернення: 10.05.2026).",
        "npm Workspaces Documentation. URL: https://docs.npmjs.com/cli/using-npm/workspaces (дата звернення: 10.05.2026).",
        "Vite Documentation. URL: https://vite.dev/ (дата звернення: 12.05.2026).",
        "React Router Documentation. URL: https://reactrouter.com/ (дата звернення: 12.05.2026).",
        "Stack Overflow Developer Survey 2024. URL: https://survey.stackoverflow.co/2024/ (дата звернення: 05.05.2026).",
        "Fowler M. Monorepo. URL: https://martinfowler.com/bliki/Monorepo.html (дата звернення: 05.05.2026).",
        "Redux Toolkit Documentation. URL: https://redux-toolkit.js.org/ (дата звернення: 10.05.2026).",
        "JSON Web Tokens. URL: https://jwt.io/ (дата звернення: 10.05.2026).",
        "SQLite Documentation. URL: https://www.sqlite.org/docs.html (дата звернення: 10.05.2026).",
    ]
    for i, ref in enumerate(references, 1):
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = FIRST_LINE_INDENT
        para.add_run(f"{i}. {ref}")


# ============================================================
# MAIN
# ============================================================

def main():
    """Generate the diploma document."""
    print("🔨 Generating diploma document...")

    doc = create_document()

    # Title page placeholder
    add_centered_paragraph(doc, "НАЦІОНАЛЬНИЙ УНІВЕРСИТЕТ «ЛЬВІВСЬКА ПОЛІТЕХНІКА»", bold=True, size=Pt(14))
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "Кафедра Системи віртуальної реальності")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "ПОЯСНЮВАЛЬНА ЗАПИСКА", bold=True, size=Pt(16))
    add_centered_paragraph(doc, "до бакалаврської кваліфікаційної роботи на тему:")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, (
        "Розроблення навчальної платформи для набуття практичних навичок "
        "full-stack розробки на основі Monorepo-архітектури"
    ), bold=True, size=Pt(14))
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, (
        "(Development of an Educational Platform for Acquiring Practical "
        "Full-Stack Development Skills Based on Monorepo Architecture)"
    ), size=Pt(12))
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "")
    add_centered_paragraph(doc, "")

    # Student info - right aligned
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.first_line_indent = Cm(0)
    para.add_run("Студентка групи КН-41д")
    para = doc.add_paragraph()
    para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.first_line_indent = Cm(0)
    run = para.add_run("Ключак О. А.")
    run.bold = True

    # Content
    write_annotation(doc)
    write_annotation_en(doc)
    write_abbreviations(doc)
    write_introduction(doc)
    write_chapter1(doc)
    write_chapter2(doc)
    write_chapter3(doc)
    write_conclusions(doc)
    write_references(doc)

    # Save
    doc.save(OUTPUT_FILE)
    print(f"✅ Document saved: {OUTPUT_FILE}")
    print(f"   Sections: Annotation, Introduction, 3 Chapters, Conclusions, References")


if __name__ == "__main__":
    main()
